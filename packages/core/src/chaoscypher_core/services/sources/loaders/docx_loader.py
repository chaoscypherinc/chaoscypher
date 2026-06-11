# Copyright (C) 2024-2026 Chaos Cypher, Inc.
# SPDX-License-Identifier: AGPL-3.0-only

"""Microsoft Word (.docx) loader.

Walks paragraphs and tables in document order. Captures heading-style
paragraphs as document landmarks (returned in metadata for citation)
and flattens table rows to ``col1 | col2 | …`` lines so plain-text
downstream consumers (extraction, embedding) keep the row association.
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from typing import TYPE_CHECKING, Any

import structlog
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from chaoscypher_core.exceptions import ValidationError
from chaoscypher_core.plugins import PluginMetadata


if TYPE_CHECKING:
    from chaoscypher_core.settings import EngineSettings
    from chaoscypher_core.utils.chunk import LocationBoundary


logger = structlog.get_logger(__name__)


class DOCXLoader:
    """Loader for Microsoft Word ``.docx`` files."""

    @property
    def supported_extensions(self) -> list[str]:
        """File extensions this loader supports."""
        return [".docx", ".DOCX"]

    @property
    def metadata(self) -> PluginMetadata:
        """Plugin metadata."""
        return PluginMetadata(
            plugin_id="docx",
            name="Word DOCX Loader",
            description="Loads Microsoft Word .docx files.",
            version="1.0.0",
            author="ChaosCypher",
            category="loader",
            builtin=True,
            origin="builtin",
            tags=["document", "office"],
        )

    def __init__(self, settings: EngineSettings | None = None) -> None:
        """Initialize DOCX loader.

        Args:
            settings: Settings instance (not currently used).
        """
        self.settings = settings

    def load_document(self, filepath: str) -> list[dict[str, Any]]:
        """Load a Word document.

        Args:
            filepath: Path to a .docx file.

        Returns:
            A single-item list with concatenated paragraph + table text
            and metadata listing the headings and counts.

        Raises:
            ValidationError: If the file is not a valid DOCX (corrupt
                OPC package, missing zip central directory, etc.).
        """
        from chaoscypher_core.services.sources.loaders.base import check_loader_file_size

        check_loader_file_size(filepath, self.settings)

        path = Path(filepath)
        try:
            doc = Document(str(path))
        except (PackageNotFoundError, KeyError, zipfile.BadZipFile) as exc:
            msg = f"File '{path.name}' is not a valid DOCX: {exc}"
            raise ValidationError(msg, field="content") from exc

        parts: list[str] = []
        headings: list[str] = []
        # Track which parts[] index each heading lives at so we can compute
        # its char offset in the final joined content for the location_index.
        heading_part_indices: list[int] = []
        paragraphs_skipped = 0

        # Body paragraphs in document order.
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                paragraphs_skipped += 1
                continue
            style_name = para.style.name.lower() if (para.style and para.style.name) else ""
            if style_name.startswith("heading") or style_name == "title":
                headings.append(text)
                heading_part_indices.append(len(parts))
                parts.append(f"\n{text}\n{'=' * len(text)}\n")
            else:
                parts.append(text)

        # Tables, flattened.
        for table_idx, table in enumerate(doc.tables):
            parts.append(f"\n[Table {table_idx + 1}]")
            for row in table.rows:
                cells = " | ".join(cell.text.strip() for cell in row.cells)
                parts.append(cells)

        content = "\n\n".join(parts)

        # Build location_index from heading offsets. parts are joined with
        # "\n\n", so parts[i]'s start in content is sum(len(parts[j]) for
        # j < i) + 2 * i. Each heading owns the char range from its own
        # start to the next heading's start (or end of content for the last).
        location_index: list[LocationBoundary] = []
        if heading_part_indices:
            part_offsets: list[int] = []
            cum = 0
            for idx, part in enumerate(parts):
                part_offsets.append(cum)
                cum += len(part)
                if idx < len(parts) - 1:
                    cum += 2  # "\n\n"

            for i, heading_part_idx in enumerate(heading_part_indices):
                start = part_offsets[heading_part_idx]
                if i + 1 < len(heading_part_indices):
                    # End is the start of the NEXT heading's part, minus the
                    # separator that belongs to the boundary itself.
                    end = part_offsets[heading_part_indices[i + 1]] - 2
                else:
                    end = len(content)
                location_index.append(
                    {
                        "start_char": start,
                        "end_char": end,
                        "page_number": None,
                        "section": headings[i],
                    }
                )

        logger.info(
            "docx_loaded",
            filepath=str(path),
            paragraph_count=len(doc.paragraphs),
            table_count=len(doc.tables),
            heading_count=len(headings),
            character_count=len(content),
            paragraphs_skipped=paragraphs_skipped,
        )

        return [
            {
                "content": content,
                "metadata": {
                    "source": str(path),
                    "extraction_method": "docx",
                    "headings": headings,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "content_type": (
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                    "loader_docx_paragraphs_skipped": paragraphs_skipped,
                    "location_index": location_index,
                },
            }
        ]

    def supports_ocr(self) -> bool:
        """DOCX files don't need OCR."""
        return False


__all__ = ["DOCXLoader"]
