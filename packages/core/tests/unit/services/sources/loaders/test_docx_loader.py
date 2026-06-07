# Copyright (C) 2024-2026 Chaos Cypher, Inc.
# SPDX-License-Identifier: AGPL-3.0-only

"""DOCXLoader extracts paragraphs, headings, and tables.

Workstream 7 (2026-05-07): standalone .docx uploads had no loader.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from chaoscypher_core.exceptions import ValidationError
from chaoscypher_core.services.sources.loaders.docx_loader import DOCXLoader


FIXTURE = Path(__file__).parent.parent.parent.parent.parent / "fixtures" / "loaders" / "sample.docx"


def test_docx_loader_extracts_paragraphs() -> None:
    assert FIXTURE.exists(), f"fixture missing: {FIXTURE}"
    loader = DOCXLoader()
    docs = loader.load_document(str(FIXTURE))
    text = docs[0]["content"]
    assert "Sample Document" in text
    assert "First paragraph" in text
    assert "Second paragraph" in text
    assert "Section content" in text
    assert docs[0]["metadata"]["extraction_method"] == "docx"


def test_docx_loader_captures_headings() -> None:
    loader = DOCXLoader()
    docs = loader.load_document(str(FIXTURE))
    headings = docs[0]["metadata"]["headings"]
    assert "Sample Document" in headings
    assert "Section" in headings


def test_docx_loader_extracts_table_rows() -> None:
    loader = DOCXLoader()
    docs = loader.load_document(str(FIXTURE))
    text = docs[0]["content"]
    # Table cells flattened to "col1 | col2" lines.
    assert "A1 | B1" in text
    assert "A2 | B2" in text
    assert docs[0]["metadata"]["table_count"] == 1


def test_docx_loader_supported_extensions() -> None:
    loader = DOCXLoader()
    assert ".docx" in loader.supported_extensions
    assert loader.metadata.plugin_id == "docx"


def test_docx_loader_wraps_invalid_file_in_validation_error(tmp_path: Path) -> None:
    """A file with a .docx extension but wrong magic bytes raises ValidationError.

    Without the wrapper, python-docx would surface either a
    ``PackageNotFoundError`` or a ``zipfile.BadZipFile`` — both bypass
    the API exception envelope and surface as 500s. With the wrapper
    it surfaces as a 422 ValidationError with the filename and a clear
    message.
    """
    bogus = tmp_path / "fake.docx"
    bogus.write_bytes(b"not a real docx")

    loader = DOCXLoader()
    with pytest.raises(ValidationError) as exc_info:
        loader.load_document(str(bogus))

    err = exc_info.value
    assert err.code == "VALIDATION_ERROR"
    assert err.field == "content"
    assert "fake.docx" in err.message
    assert "not a valid" in err.message
    assert exc_info.value.__cause__ is not None


# ---------------------------------------------------------------------------
# 2026-05-18: location_index emitted from heading-style paragraphs
# ---------------------------------------------------------------------------


def test_docx_loader_emits_location_index_with_headings() -> None:
    """DOCX loader emits a location_index entry for each heading-style
    paragraph. section is the heading text; page_number is None.
    """
    loader = DOCXLoader()
    docs = loader.load_document(str(FIXTURE))
    metadata = docs[0]["metadata"]
    headings = metadata["headings"]

    location_index = metadata["location_index"]
    assert len(location_index) == len(headings)

    for entry, heading in zip(location_index, headings, strict=True):
        assert entry["page_number"] is None
        assert entry["section"] == heading


def test_docx_loader_emits_empty_location_index_when_no_headings(tmp_path: Path) -> None:
    """DOCX with no heading-style paragraphs emits an empty location_index."""
    from docx import Document as _Document

    doc = _Document()
    doc.add_paragraph("Just a regular paragraph.")
    doc.add_paragraph("Another regular paragraph.")
    path = tmp_path / "no_headings.docx"
    doc.save(str(path))

    loader = DOCXLoader()
    docs = loader.load_document(str(path))
    metadata = docs[0]["metadata"]

    assert metadata["headings"] == []
    assert metadata["location_index"] == []
